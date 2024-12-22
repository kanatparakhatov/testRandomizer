import random
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches



def clean_and_load_questions(input_path):
    """Читает, очищает и загружает вопросы с ответами из файла."""
    with open(input_path, 'r', encoding='utf-8') as file:
        lines = [line.strip() for line in file if line.strip()]  # Убираем пустые строки

    # Объединяем строки в один текст и разделяем на блоки вопросов
    content = '\n'.join(lines)
    questions = content.split('+++++')

    parsed_questions = []
    for question_block in questions:
        lines = question_block.strip().split('\n')
        if len(lines) > 1:
            question = lines[0].strip()
            answers = [line.strip().lstrip('=') for line in lines[1:] if line.strip() and not line.startswith('====')]
            # Если ответов 5, первый — это продолжение вопроса
            if len(answers) == 5:
                question += f" {answers[0]}"
                answers = answers[1:]
            parsed_questions.append((question, answers))
    return parsed_questions


def set_default_style(paragraph):
    """Устанавливает стандартный стиль текста: шрифт Times New Roman, размер 11, черный цвет."""
    # Устанавливаем стандартный стиль для всего текста абзаца
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False
        run.font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет


def remove_extra_spacing(paragraph):
    """Удаляет дополнительные отступы и интервалы."""
    # Убираем межстрочные интервалы
    paragraph.paragraph_format.line_spacing = Pt(12)  # Одинарный интервал
    paragraph.paragraph_format.space_after = Pt(0)  # Убираем интервал после абзаца
    paragraph.paragraph_format.space_before = Pt(0)  # Убираем интервал перед абзацем


def create_test_variants(questions, num_questions, num_variants, output_dir, subject):
    """Создаёт файлы с уникальными вариантами тестов в формате Word."""
    if len(questions) < num_questions:
        raise ValueError("Недостаточно вопросов в исходном файле для создания варианта.")

    for variant in range(1, num_variants + 1):
        # Выбираем случайные вопросы
        variant_questions = random.sample(questions, num_questions)

        # Создание нового документа Word
        doc = Document()

        # Установка полей документа
        section = doc.sections[0]
        section.left_margin = Inches(1.18)  # 3 см
        section.right_margin = Inches(0.59)  # 1.5 см
        section.top_margin = Inches(0.79)  # 2 см
        section.bottom_margin = Inches(0.79)  # 2 см

        # Добавляем заголовок, применяя тот же стиль
        heading = doc.add_paragraph()
        heading.add_run(
            f"_____________  guruh talabasi ___________________________________________________  \"{subject}\" fan farqini topshirish bo'yicha nazorat ishi.")
        set_default_style(heading)

        # Убираем лишние отступы после заголовка
        remove_extra_spacing(heading)

        # Добавляем строку с номером варианта и выравниваем её по центру
        variant_paragraph = doc.add_paragraph()
        run_variant = variant_paragraph.add_run(f"Variant №{variant}")
        run_variant.bold = True
        set_default_style(variant_paragraph)

        # Устанавливаем выравнивание по центру для строки "Variant №X"
        variant_paragraph.alignment = 1  # Выравнивание по центру

        # Убираем лишние отступы после строки с номером варианта
        remove_extra_spacing(variant_paragraph)

        # Добавляем отступ 5 пт после строки с номером варианта
        paragraph = doc.add_paragraph()  # Пустой абзац для отступа
        paragraph.paragraph_format.space_before = Pt(5)  # Отступ в 5 пунктов
        remove_extra_spacing(paragraph)

        # Добавляем вопросы и ответы
        for i, (q, ans) in enumerate(variant_questions):
            # Добавляем отступ перед вопросом
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{i + 1}. {q}")  # Вставляем номер вопроса
            set_default_style(paragraph)  # Применяем стиль
            paragraph.paragraph_format.space_before = Pt(5)  # Добавляем отступ перед вопросом

            # Убираем лишние отступы после вопроса
            remove_extra_spacing(paragraph)

            # Добавляем ответы с табуляцией между вопросом и ответом
            for j, a in enumerate(ans):
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"\t{chr(97 + j)}. {a}")  # Используем табуляцию перед ответом
                set_default_style(paragraph)  # Применяем стиль

                # Убираем лишние отступы после ответа
                remove_extra_spacing(paragraph)

            # Добавляем отступ после последнего ответа
            paragraph = doc.add_paragraph()  # Пустой абзац для отступа
            paragraph.paragraph_format.space_before = Pt(5)  # Отступ в 5 пунктов

            # Убираем лишние отступы после пустого абзаца
            remove_extra_spacing(paragraph)

        # Сохраняем документ
        output_path = f"{output_dir}/variant_{variant}.docx"
        doc.save(output_path)
        print(f"Создан файл: {output_path}")


if __name__ == "__main__":
    # Пути к файлам
    input_path = "C:/asd/test/uzb_his_uzb.txt"
    output_directory = "c:/asd/finishedTest/uzb_his_uzb/"

    # Очищаем и загружаем вопросы
    print("Очистка и загрузка вопросов из файла...")
    questions_data = clean_and_load_questions(input_path)

    # Запрос параметров для генерации тестов
    subject = input("Введите предмет: ")
    num_questions = int(input("Введите количество вопросов в одном варианте: "))
    num_variants = int(input("Введите количество вариантов: "))

    # Создание тестовых вариантов
    os.makedirs(output_directory, exist_ok=True)
    create_test_variants(questions_data, num_questions, num_variants, output_directory, subject)