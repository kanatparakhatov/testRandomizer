import random
from docx import Document

# Функция для загрузки вопросов из файла docx
def load_questions_from_docx(file_path):
    """Загружает вопросы из указанного .docx файла без нумерации."""
    document = Document(file_path)
    questions = [para.text.strip().lstrip('0123456789. ').strip() for para in document.paragraphs if para.text.strip()]
    return questions

# Функция для создания вариантов
def create_variants(questions, num_variants, questions_per_variant):
    """Создает список вариантов с указанным количеством вопросов без повторений."""
    if len(questions) < num_variants * questions_per_variant:
        raise ValueError("Недостаточно вопросов для создания всех вариантов без повторений.")

    random.shuffle(questions)  # Перемешиваем вопросы
    variants = []

    for i in range(num_variants):
        start_index = i * questions_per_variant
        end_index = start_index + questions_per_variant
        variant = questions[start_index:end_index]
        variants.append(variant)

    return variants

# Функция для сохранения вариантов в новый docx файл
def save_variants_to_docx(variants, output_path):
    """Сохраняет варианты вопросов в новый .docx файл."""
    document = Document()
    for i, variant in enumerate(variants, 1):
        document.add_heading(f'Вариант {i}', level=1)
        for j, question in enumerate(variant, 1):
            document.add_paragraph(f'{j}. {question}')
        document.add_paragraph()  # Пустая строка между вариантами
    document.save(output_path)

# Основной код
if __name__ == "__main__":
    input_file = 'c:/1.docx'  # Укажите путь к вашему исходному файлу
    output_file = 'c:/random_variants1.docx'  # Укажите путь для сохранения нового файла

    try:
        # Загружаем вопросы
        questions = load_questions_from_docx(input_file)

        if len(questions) < 5:
            print("Недостаточно вопросов в файле! Нужно минимум 5 вопросов.")
        else:
            # Создаем варианты
            num_variants = 30  # Количество вариантов
            questions_per_variant = 5  # Количество вопросов в каждом варианте

            variants = create_variants(questions, num_variants, questions_per_variant)

            # Сохраняем варианты в файл
            save_variants_to_docx(variants, output_file)
            print(f"Случайные варианты успешно сохранены в файл: {output_file}")
    except Exception as e:
        print(f"Ошибка: {e}")
